import openpyxl
import docx
import csv
import os
import exp_finder as f

# constants 
cwd = os.getcwd()
input_directory = 'py-files\\data\\testdc'
input_file_name = 'fulltestdc.xlsx'
input_path = os.path.join(cwd, input_directory, input_file_name)
output_directory = 'py-files\\data\\resultsdc'
output_file_name = 'results.docx'
output_path = os.path.join(cwd, output_directory, output_file_name)
output_analysis_file_name = 'results.csv'
output_analysis_directory = 'py-files\\data\\resultsdc'
output_analysis_path = os.path.join(cwd, output_analysis_directory, output_analysis_file_name)

analysis = {
  'norms': set(),
  'appendixes': set(),
  'acronyms': set(),
  'count_lsc4_chars': 0,
  'count_lsc2_chars': 0,
  'count_lsc4_words': 0,
  'count_lsc2_words': 0
}
words_per_page = 363.25

wb = openpyxl.load_workbook(input_path)
print('-> Excel source file opened successfully')
tabs = wb.sheetnames

letters = [ 'A', 'B', 'C', 'D', 'E', 'F']
doc = docx.Document(output_path)
print('-> Word target file opened successfully')
styles = doc.styles
i = 0

for tab in tabs:
  print(' > transfering worksheet ', tab)
  # doc.add_heading(tab, level=2)
  ws = wb[tab]
  max_rows = ws.max_row

  heading = doc.add_paragraph(tab.upper())
  heading.style = styles['Normal']
  table = doc.add_table(rows=0, cols=3)
  table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
  previous_code = ''
  previous_title = ''

  for i in range(2, max_rows + 1):
    position_code = 'A' + str(i)
    code = str(ws[position_code].value).strip()
    if (code == 'None'): continue
    position_title = 'B' + str(i)
    title = str(ws[position_title].value).strip()
    position_text = 'F' + str(i)
    text = str(ws[position_text].value).strip()
    position_lsc2 = 'C' + str(i)
    lsc2_text = str(ws[position_lsc2].value)
    if (previous_code != code and code != 'None'): 
      previous_code = code
      code_changed = True
    else:
      code = previous_code
      code_changed = False
    if (previous_title != title and title != 'None'): 
      previous_title = title
    else:
      title = ''
    if (title != 'None' and text != 'None'):
      cells = table.add_row().cells
      if (code != 'None' and code_changed): 
        cells[0].text = code
      if (title != 'None'): 
        cells[1].text = title
      if (text != 'None'): 
        cells[2].text = text
      else:
        cells[2].text = ''
    i += 1
    acronyms = f.find_acronyms(text)
    for acronym in acronyms:
      analysis['acronyms'].add(acronym)
    appendixes = f.find_appendixes(text)
    for appendix in appendixes:
      analysis['appendixes'].add(appendix)
    norms = f.find_norms(text)
    for norm in norms:
      analysis['norms'].add(norm)
    analysis['count_lsc2_chars'] += len(lsc2_text)
    analysis['count_lsc4_chars'] += len(text)
    analysis['count_lsc2_words'] += len(lsc2_text.split())
    analysis['count_lsc4_words'] += len(text.split())

  par = doc.add_paragraph()
  run = par.add_run()
  run.add_break(docx.enum.text.WD_BREAK.PAGE)

analysis['acronyms'] = sorted(analysis['acronyms'])
analysis['appendixes'] = sorted(analysis['appendixes'])
analysis['norms'] = sorted(analysis['norms'])

#end
doc.save(output_path)
print('-> Word target file saved successfully')
print('-> Transfer completed')
print('-> Sending analytical results to csv file')
with open(output_analysis_path, 'w', newline='', encoding='utf-8') as file_output:
  writer = csv.writer(file_output, delimiter=' ')
  writer.writerow(['\n\nWORDCOUNT LSC2'])
  writer.writerow([str(analysis['count_lsc2_chars']), 'characters'])
  writer.writerow([str(analysis['count_lsc2_words']), 'words'])
  pages_lsc2 = round(analysis['count_lsc2_words'] / words_per_page, 2)
  writer.writerow([str(pages_lsc2), 'pages'])
  writer.writerow(['\n\nWORDCOUNT LSC4'])
  writer.writerow([str(analysis['count_lsc4_chars']), 'characters'])
  writer.writerow([str(analysis['count_lsc4_words']), 'words'])    
  pages_lsc4 = round(analysis['count_lsc4_words'] / words_per_page, 2)
  writer.writerow([str(pages_lsc4), 'pages'])
  writer.writerow(['\n\nACRONYMS'])
  for row in analysis['acronyms']:
    writer.writerow([str(row)])
  writer.writerow(['\n\nNORMS'])
  for row in analysis['norms']:
    writer.writerow([str(row)])
  writer.writerow(['\n\nAPPENDIXES'])
  for row in analysis['appendixes']:
    writer.writerow([str(row)])

  file_output.close()
print('-> results.csv file created successfully')